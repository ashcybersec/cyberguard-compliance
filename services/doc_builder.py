import os
import uuid
import zipfile
import re
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


# ── Brand colours ──────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x2A)
CYAN   = RGBColor(0x00, 0xC2, 0xCB)
GREY   = RGBColor(0x55, 0x66, 0x77)
LGREY  = RGBColor(0x88, 0x99, 0xAA)
BLACK  = RGBColor(0x1A, 0x1A, 0x2E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

DOCUMENT_TITLES = {
    "scope_statement":             "Cyber Essentials — Scope Statement",
    "information_security_policy": "Information Security Policy",
    "access_control_policy":       "Access Control Policy",
    "patch_management_policy":     "Patch Management Policy",
    "firewall_network_policy":     "Firewall & Network Security Policy",
    "malware_protection_policy":   "Malware Protection Policy",
    "asset_register_guidance":     "Asset Register — Guidance & Template",
}

DOCUMENT_REFS = {
    "scope_statement":             "CG-SCOPE-001",
    "information_security_policy": "CG-POL-001",
    "access_control_policy":       "CG-POL-002",
    "patch_management_policy":     "CG-POL-003",
    "firewall_network_policy":     "CG-POL-004",
    "malware_protection_policy":   "CG-POL-005",
    "asset_register_guidance":     "CG-REG-001",
}


class EvidencePackBuilder:
    OUTPUT_DIR = os.getenv("REPORT_OUTPUT_DIR", "/tmp/reports")

    def __init__(self):
        os.makedirs(self.OUTPUT_DIR, exist_ok=True)

    def build_zip(self, company_name: str, documents: dict) -> str:
        session_id = uuid.uuid4().hex[:8]
        safe_name  = company_name.replace(" ", "_").replace("/", "-")[:40]
        zip_path   = f"{self.OUTPUT_DIR}/CyberGuard_EvidencePack_{safe_name}_{session_id}.zip"

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for i, (doc_key, content) in enumerate(documents.items(), 1):
                title    = DOCUMENT_TITLES.get(doc_key, doc_key.replace("_", " ").title())
                ref      = DOCUMENT_REFS.get(doc_key, f"CG-DOC-{i:03d}")
                filename = f"{i:02d}_{doc_key}.docx"
                docx_bytes = self._build_docx(company_name, title, ref, content)
                zf.writestr(filename, docx_bytes)

        return zip_path

    def _build_docx(self, company: str, title: str, ref: str, content: str) -> bytes:
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Cm(2.2)
            section.bottom_margin = Cm(2.2)
            section.left_margin   = Cm(2.8)
            section.right_margin  = Cm(2.5)

        # Header
        header = doc.sections[0].header
        htbl   = header.add_table(1, 2, width=Inches(6.5))
        self._remove_table_borders(htbl)

        left  = htbl.cell(0, 0).paragraphs[0]
        right = htbl.cell(0, 1).paragraphs[0]

        rl = left.add_run("CyberGuard Essentials")
        rl.bold = True
        rl.font.size = Pt(10)
        rl.font.color.rgb = CYAN
        rl.font.name = "Calibri"
        left.alignment = WD_ALIGN_PARAGRAPH.LEFT

        rr = right.add_run(f"{ref}  ·  {company}")
        rr.font.size = Pt(8)
        rr.font.color.rgb = LGREY
        rr.font.name = "Calibri"
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        self._add_cyan_rule(header)

        # Cover block
        doc.add_paragraph()

        pc = doc.add_paragraph()
        rc = pc.add_run(company.upper())
        rc.font.size = Pt(8)
        rc.font.color.rgb = LGREY
        rc.font.name = "Calibri"
        rc.bold = True
        pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pc.paragraph_format.space_after = Pt(4)

        pt = doc.add_paragraph()
        rt = pt.add_run(title)
        rt.bold = True
        rt.font.size = Pt(22)
        rt.font.color.rgb = NAVY
        rt.font.name = "Calibri"
        pt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pt.paragraph_format.space_after = Pt(2)

        # Thick cyan rule under title
        self._add_thick_rule(doc, "00C2CB", thickness=18)

        pm = doc.add_paragraph()
        rm = pm.add_run(
            f"Version 1.0  ·  {date.today().strftime('%d %B %Y')}  ·  "
            f"NCSC Cyber Essentials v3.3  ·  CONFIDENTIAL"
        )
        rm.font.size = Pt(8.5)
        rm.font.color.rgb = LGREY
        rm.font.name = "Calibri"
        pm.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pm.paragraph_format.space_after = Pt(16)

        # Body
        self._render_content(doc, content)

        # Footer
        footer = doc.sections[0].footer
        self._add_cyan_rule(footer)
        fp = footer.add_paragraph()
        fp.add_run(
            f"© {date.today().year} Corvaxis Ltd · ceready.co.uk  ·  "
            f"NCSC Cyber Essentials v3.3  ·  {ref}"
        ).font.size = Pt(7.5)
        fp.runs[0].font.color.rgb = LGREY
        fp.runs[0].font.name = "Calibri"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _clean_content(self, content: str) -> str:
        """Clean and normalise content before rendering."""
        # Unescape JSON string escapes
        content = content.replace('\\n', '\n')
        content = content.replace('\\"', '"')
        content = content.replace("\\'", "'")
        content = content.replace('\\t', '    ')

        # Strip JSON wrapper if present
        stripped = content.strip()
        if stripped.startswith('{'):
            try:
                import json
                parsed = json.loads(stripped)
                content = list(parsed.values())[0] if parsed else content
            except Exception:
                pass

        # Remove lines that are just underscores (horizontal rules) — we'll use proper Word rules
        lines = content.split('\n')
        cleaned = []
        for line in lines:
            stripped_line = line.strip()
            # Skip lines that are just underscores, dashes or equals (markdown dividers)
            if re.match(r'^[_\-=]{4,}$', stripped_line):
                continue
            cleaned.append(line)

        return '\n'.join(cleaned)

    def _render_content(self, doc: Document, content: str):
        content = self._clean_content(content)
        lines = content.split('\n')
        i = 0
        prev_was_heading = False

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Empty line
            if not stripped:
                i += 1
                continue

            # H1
            if stripped.startswith('# ') and not stripped.startswith('##'):
                p = doc.add_heading(stripped[2:].strip(), level=1)
                self._style_h1(p)
                prev_was_heading = True
                i += 1
                continue

            # H2
            if stripped.startswith('## '):
                if not prev_was_heading:
                    sp = doc.add_paragraph()
                    sp.paragraph_format.space_before = Pt(12)
                    sp.paragraph_format.space_after  = Pt(0)
                p = doc.add_heading(stripped[3:].strip(), level=2)
                self._style_h2(p)
                prev_was_heading = True
                i += 1
                continue

            # H3
            if stripped.startswith('### '):
                p = doc.add_heading(stripped[4:].strip(), level=3)
                self._style_h3(p)
                prev_was_heading = True
                i += 1
                continue

            prev_was_heading = False

            # Table
            if stripped.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._render_table(doc, table_lines)
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                continue

            # Bullet
            if stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:].strip()
                text = self._strip_bold(text)
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_run(p, text, size=10)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(3)
                i += 1
                continue

            # Numbered list
            if re.match(r'^\d+[\.\)]\s', stripped):
                text = re.sub(r'^\d+[\.\)]\s+', '', stripped)
                text = self._strip_bold(text)
                p = doc.add_paragraph(style='List Number')
                self._add_formatted_run(p, text, size=10)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(3)
                i += 1
                continue

            # Standalone bold line = subheading
            if stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
                text = stripped.strip('*').strip()
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(10.5)
                r.font.color.rgb = NAVY
                r.font.name = "Calibri"
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(3)
                i += 1
                continue

            # Regular paragraph
            text = self._strip_bold(stripped)
            p = doc.add_paragraph()
            self._add_formatted_run(p, text, size=10)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.first_line_indent = Pt(0)
            i += 1

    def _strip_bold(self, text: str) -> str:
        """Remove markdown bold markers from text."""
        return re.sub(r'\*\*(.*?)\*\*', r'\1', text)

    def _add_formatted_run(self, p, text: str, size: float = 10, bold: bool = False, colour: RGBColor = None):
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        r.bold = bold
        if colour:
            r.font.color.rgb = colour
        else:
            r.font.color.rgb = BLACK

    def _render_table(self, doc: Document, table_lines: list):
        rows = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip('|').split('|')]
            # Skip separator rows
            if all(re.match(r'^[-: ]+$', c) for c in cells if c):
                continue
            # Clean bold markers from cells
            cells = [self._strip_bold(c) for c in cells]
            rows.append(cells)

        if not rows:
            return

        ncols = max(len(r) for r in rows)
        tbl   = doc.add_table(rows=len(rows), cols=ncols)
        tbl.style = 'Table Grid'

        for ri, row_data in enumerate(rows):
            for ci in range(ncols):
                cell_text = row_data[ci] if ci < len(row_data) else ''
                cell = tbl.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                r = p.add_run(cell_text)
                r.font.size = Pt(9)
                r.font.name = "Calibri"

                if ri == 0:
                    r.bold = True
                    r.font.color.rgb = WHITE
                    self._set_cell_bg(cell, '0D1B2A')
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    r.font.color.rgb = BLACK
                    if ri % 2 == 0:
                        self._set_cell_bg(cell, 'F0F4F8')

                cell.paragraphs[0].paragraph_format.space_before = Pt(3)
                cell.paragraphs[0].paragraph_format.space_after  = Pt(3)

    def _style_h1(self, p):
        for run in p.runs:
            run.font.name  = "Calibri"
            run.font.size  = Pt(16)
            run.font.color.rgb = NAVY
            run.bold = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        self._add_para_bottom_border(p, "00C2CB", sz=6)

    def _style_h2(self, p):
        for run in p.runs:
            run.font.name  = "Calibri"
            run.font.size  = Pt(13)
            run.font.color.rgb = NAVY
            run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)

    def _style_h3(self, p):
        for run in p.runs:
            run.font.name  = "Calibri"
            run.font.size  = Pt(11)
            run.font.color.rgb = CYAN
            run.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

    def _add_para_bottom_border(self, p, hex_color: str, sz: int = 6):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(sz))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), hex_color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_cyan_rule(self, parent):
        p = parent.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '00C2CB')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_thick_rule(self, doc: Document, hex_color: str, thickness: int = 18):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(8)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(thickness))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), hex_color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _set_cell_bg(self, cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _remove_table_borders(self, table):
        tbl   = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for name in ['top','left','bottom','right','insideH','insideV']:
            b = OxmlElement(f'w:{name}')
            b.set(qn('w:val'), 'none')
            tblBorders.append(b)
        tblPr.append(tblBorders)
